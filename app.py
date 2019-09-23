from flask import Flask, render_template, request, redirect, url_for
from werkzeug.utils import secure_filename # 使用这个是为了确保filename是安全的
from os import path
import docx
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate, InlineImage
# for height and width you have to use millimeters (Mm), inches or points(Pt) class :
from docx.shared import Mm, Inches, Pt
import jinja2
app = Flask(__name__)


@app.route('/')
def hello_world():
    return render_template('/index.html')


@app.route("/upload", methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        print("eqweq")
        # f = request.files["file"]
        f = request.files['file']
        print("dada")
        # base_path = path.abspath(path.dirname(__file__))
        # upload_path = path.join(base_path, 'static/uploads/')
        # file_name = upload_path + secure_filename(f.filename)
        document = Document(f)
        for paragraph in document.paragraphs:
            # print(paragraph.text)
            for run in paragraph.runs:
                  print(run.text)
                  run.text = run.replace('一','一')
                  run.text = run.text.replace('、', '{{myimage}}   ')
                  run.text = run.replace('一', '一')
        # for i in range(len(document.paragraphs)):
        #     print("第" + str(i) + "段的内容是：" + document.paragraphs[i].text)
        #     word = document.paragraphs[i].text.split('、')
        #     print(word)
        #     for i in word:
        #         j =  word.index(i)
        #         word.insert(j+1,'{{ myimage }}')
        #     print(word)
          # print(text)

        document.save('demo.docx')
        tpl = DocxTemplate('demo.docx')
        context = {
            'myimage': InlineImage(tpl, '标号.png', width=Mm(8)),
        }
        # testing that it works also when autoescape has been forced to True
        jinja_env = jinja2.Environment(autoescape=True)
        tpl.render(context, jinja_env)
        tpl.save('end.docx')
    return render_template('upload.html')


if __name__ == '__main__':
    app.run()
